"""测试 docx → PDF 转换完整链路"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.agents.tools.python_executor_tool import PythonExecutorTool
from pathlib import Path

tool = PythonExecutorTool()
PROJECT_ROOT = Path(__file__).resolve().parents[1]

# 先创建一个简单的测试 docx（纯 ASCII 文件名，避免 COM 中文名问题）
print("=== Step 1: 创建测试 docx ===")
code_create = """
from docx import Document
import os

doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph for docx to PDF conversion.')
doc.add_paragraph('JChatMind Python version test.')
doc.add_heading('Section 2', level=1)
doc.add_paragraph('Chinese: 中文内容测试 Testing Chinese content.')
doc_path = os.path.join(GENERATED_DIR, 'test_input.docx')
doc.save(doc_path)
print('Created:', doc_path)
"""
r = tool._run(code=code_create)
print(r)

# 使用绝对路径转换
generated_dir = PROJECT_ROOT / "uploads" / "generated"
test_docx = generated_dir / "test_input.docx"

if test_docx.exists():
    print(f"\n=== Step 2: 转换 docx → PDF ===")
    abs_src = str(test_docx).replace("\\", "/")
    code_convert = f"""
from docx2pdf import convert
import os

src = r'{abs_src}'
out = OUTPUT_FILE

print('Converting:', src)
print('Output:', out)
convert(src, out)
if os.path.exists(out):
    print('PDF size:', os.path.getsize(out), 'bytes')
    print('SUCCESS')
else:
    print('FAILED - output not found')
"""
    result = tool._run(code=code_convert, output_filename="test_output.pdf")
    print(result)
else:
    print("Step 1 failed - test_input.docx not created")
